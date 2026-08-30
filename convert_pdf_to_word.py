"""Convert scanned PDF to editable Word document using OCR."""

from __future__ import annotations

import sys
import tempfile
from pathlib import Path

import pypdfium2 as pdfium
import pytesseract
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image, ImageEnhance, ImageFilter, ImageOps

TESSERACT_CMD = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
TESSDATA_DIR = Path(__file__).resolve().parent / "tessdata"
LANG = "por+eng"
DPI = 300
TESSERACT_CONFIG = f"--tessdata-dir {TESSDATA_DIR} --psm 3 --oem 1"


def group_words_into_lines(words: list[dict]) -> list[list[dict]]:
    if not words:
        return []

    words = sorted(words, key=lambda w: (w["top"], w["left"]))
    lines: list[list[dict]] = []
    current_line: list[dict] = []
    current_top = words[0]["top"]

    for word in words:
        if abs(word["top"] - current_top) <= 12:
            current_line.append(word)
        else:
            if current_line:
                lines.append(sorted(current_line, key=lambda w: w["left"]))
            current_line = [word]
            current_top = word["top"]

    if current_line:
        lines.append(sorted(current_line, key=lambda w: w["left"]))

    return lines


def line_text(line: list[dict]) -> str:
    return " ".join(word["text"].strip() for word in line if word["text"].strip())


def line_height(line: list[dict]) -> float:
    heights = [word["height"] for word in line if word["height"] > 0]
    return sum(heights) / len(heights) if heights else 0


def is_centered(line: list[dict], page_width: int, tolerance: float = 0.12) -> bool:
    left = min(word["left"] for word in line)
    right = max(word["left"] + word["width"] for word in line)
    line_center = (left + right) / 2
    page_center = page_width / 2
    return abs(line_center - page_center) <= page_width * tolerance


def save_temp_image(image: Image.Image, page_index: int) -> Path:
    temp_dir = Path(tempfile.gettempdir()) / "contrato_aluguel_ocr"
    temp_dir.mkdir(parents=True, exist_ok=True)
    image_path = temp_dir / f"page_{page_index + 1}.png"
    image.save(image_path, format="PNG")
    return image_path


def preprocess_image(image: Image.Image) -> Image.Image:
    gray = ImageOps.grayscale(image)
    enhanced = ImageEnhance.Contrast(gray).enhance(1.8)
    sharpened = enhanced.filter(ImageFilter.SHARPEN)
    return sharpened


def extract_words(image: Image.Image) -> list[dict]:
    processed = preprocess_image(image)
    data = pytesseract.image_to_data(
        processed,
        lang=LANG,
        config=TESSERACT_CONFIG,
        output_type=pytesseract.Output.DICT,
    )
    words: list[dict] = []

    for i, text in enumerate(data["text"]):
        cleaned = (text or "").strip()
        if not cleaned:
            continue
        conf = int(float(data["conf"][i])) if data["conf"][i] != "-1" else -1
        if conf < 35:
            continue
        words.append(
            {
                "text": cleaned,
                "left": int(data["left"][i]),
                "top": int(data["top"][i]),
                "width": int(data["width"][i]),
                "height": int(data["height"][i]),
            }
        )

    return words


def add_line_to_doc(document: Document, text: str, *, centered: bool, height: float, body_height: float) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"

    if centered and height >= body_height * 1.15:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.bold = True
        run.font.size = Pt(14)
    elif centered:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.font.size = Pt(12)
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run.font.size = Pt(12)

    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15


def convert_pdf_to_docx(pdf_path: Path, docx_path: Path) -> None:
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD

    pdf = pdfium.PdfDocument(str(pdf_path))
    document = Document()

    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("CONTRATO DE ALUGUEL 2027")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = "Times New Roman"

    subtitle = document.add_paragraph("Versão editável em Word")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(11)
    subtitle.runs[0].italic = True

    document.add_paragraph()

    scale = DPI / 72
    page_images: list[Image.Image] = []

    for page_index in range(len(pdf)):
        page = pdf[page_index]
        image = page.render(scale=scale).to_pil()
        page_images.append(image)

        ref = document.add_paragraph()
        ref.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ref.add_run().add_picture(str(save_temp_image(image, page_index)), width=Inches(6.5))

        if page_index < len(pdf) - 1:
            document.add_page_break()

    document.add_page_break()
    heading = document.add_paragraph("Texto para edição")
    heading.runs[0].bold = True
    heading.runs[0].font.size = Pt(14)

    note = document.add_paragraph(
        "O texto abaixo foi extraído automaticamente do documento digitalizado. "
        "Revise nomes, datas e valores com base na imagem original acima."
    )
    note.runs[0].italic = True
    note.runs[0].font.size = Pt(10)

    document.add_paragraph()

    for page_index, image in enumerate(page_images):

        words = extract_words(image)
        lines = group_words_into_lines(words)

        if not lines:
            continue

        heights = [line_height(line) for line in lines if line_height(line) > 0]
        body_height = sorted(heights)[len(heights) // 2] if heights else 12

        previous_top = None
        for line in lines:
            text = line_text(line)
            if not text:
                continue

            top = min(word["top"] for word in line)
            if previous_top is not None and top - previous_top > body_height * 1.8:
                document.add_paragraph()

            add_line_to_doc(
                document,
                text,
                centered=is_centered(line, image.width),
                height=line_height(line),
                body_height=body_height,
            )
            previous_top = top

        if page_index < len(pdf) - 1:
            document.add_page_break()

    document.save(docx_path)


def main() -> int:
    pdf_path = Path(r"c:\Users\Pessoal\Downloads\CONTRATO DE ALUGUEL 2027.pdf")
    docx_path = Path(r"c:\Users\Pessoal\Downloads\CONTRATO DE ALUGUEL 2027.docx")

    if not pdf_path.exists():
        print(f"Arquivo nao encontrado: {pdf_path}")
        return 1

    print(f"Convertendo: {pdf_path}")
    convert_pdf_to_docx(pdf_path, docx_path)
    print(f"Arquivo gerado: {docx_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
